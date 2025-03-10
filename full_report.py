import psycopg2
import tkinter as tk
from tkinter import messagebox, scrolledtext, filedialog
from docx import Document


# Example columns list (fill yours later)
columns_to_check = ["c_pano_av", "syno", "pht_mas_a", "pht_mas_b", "pht_mas_c", "pht_mas_d", "ch_fer_apr", "c_ouv_ap2", "c_pano_apr", "pho_fer_av", "c_ouv_av_1"]

# Define full names for the columns
column_full_names = {
    "c_pano_av": "Chambre Panoramique Avant",
    "syno": "Photo Feuille Manuel",
    "pht_mas_a": "Photo Masque A",
    "pht_mas_b":"Photo Masque B",
    "pht_mas_c":"Photo Masque C",
    "pht_mas_d":"Photo Masque D",
    "ch_fer_apr":"Chambre Fermeture Apres",
    "c_ouv_ap2":"Chambre Ouverte Après Photo",
    "c_pano_apr":"Chambre Panoramique Apres",
    "pho_fer_av":"Photo Fermeture Avant",
    "c_ouv_av_1":"Chambre Ouverte Avant",
}

# Modified SQL Queries Dictionary
SQL_QUERIES = {
    "total_count": """
        SELECT COUNT({col})
        FROM {table};
    """,
    "file_not_found_count": """
        SELECT COUNT({col})
        FROM {table}
        WHERE {col} ILIKE 'File Not Found%';
    """,
    "link_not_found_count": """
        SELECT COUNT({col})
        FROM {table}
        WHERE {col} ILIKE 'Link Not Found%';
    """,
    "file_not_found_ids": """
        SELECT id, id_troncon, code
        FROM {table}
        WHERE {col} ILIKE 'File Not Found%';
    """,
    "link_not_found_ids": """
        SELECT id, id_troncon, code
        FROM {table}
        WHERE {col} ILIKE 'Link Not Found%';
    """
}

def gather_column_info(cursor, table_name, column_name):
    column_info = {}
    # Get counts
    for check_name in ["total_count", "file_not_found_count", "link_not_found_count"]:
        cursor.execute(SQL_QUERIES[check_name].format(col=column_name, table=table_name))
        column_info[check_name] = cursor.fetchone()[0]
    
    # Get IDs for "File Not Found"
    cursor.execute(SQL_QUERIES["file_not_found_ids"].format(col=column_name, table=table_name))
    column_info["file_not_found_ids"] = cursor.fetchall()
    
    # Get IDs for "Link Not Found"
    cursor.execute(SQL_QUERIES["link_not_found_ids"].format(col=column_name, table=table_name))
    column_info["link_not_found_ids"] = cursor.fetchall()
    
    return column_info

def generate_report(conn, table_name):
    report = {}
    with conn.cursor() as cursor:
        for col in columns_to_check:
            report[col] = gather_column_info(cursor, table_name, col)
    return report

def display_report_gui(report):
    window = tk.Tk()
    window.title("File and Link Status Report")
    window.geometry("800x600")

    # Create a scrolled text area
    text_area = scrolledtext.ScrolledText(window, wrap=tk.WORD, width=95, height=30, font=("Arial", 10))
    text_area.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

    # Define tags for formatting
    text_area.tag_configure("bold", font=("Arial", 10, "bold"))
    text_area.tag_configure("title", font=("Arial", 12, "bold"))
    text_area.tag_configure("center", justify="center")
    text_area.tag_configure("spacing", spacing1=5, spacing3=5)

    # Add a title
    text_area.insert(tk.END, "FILE AND LINK STATUS REPORT\n\n", ("title", "center", "spacing"))
    
    for column, data in report.items():
        # Add column name (centered and bold)
        full_name = column_full_names.get(column, column)
        text_area.insert(tk.END, f"Column: {column} ({full_name})\n", ("bold", "center", "spacing"))
        
        # Add counts
        text_area.insert(tk.END, f"  Total Count: {data['total_count']}\n", "spacing")
        text_area.insert(tk.END, f"  'File Not Found' Count: {data['file_not_found_count']}\n", "spacing")
        text_area.insert(tk.END, f"  'Link Not Found' Count: {data['link_not_found_count']}\n", "spacing")
        
        # Add IDs for "File Not Found"
        if data['file_not_found_ids']:
            text_area.insert(tk.END, "\n  IDs with 'File Not Found':\n", "bold")
            for row in data['file_not_found_ids']:
                text_area.insert(tk.END, f"ID: {row[0]}, ID Tronc: {row[1]}, Code: {row[2]} \n")
        
        # Add IDs for "Link Not Found"
        if data['link_not_found_ids']:
            text_area.insert(tk.END, "\n  IDs with 'Link Not Found':\n", "bold")
            for row in data['link_not_found_ids']:
                text_area.insert(tk.END, f"ID: {row[0]}, ID Tronc: {row[1]}, Code: {row[2]} \n")
                
        text_area.insert(tk.END, "\n" + "-" * 80 + "\n\n", "spacing")

    # Disable editing
    text_area.config(state=tk.DISABLED)

    # Add a save button
    save_button = tk.Button(window, text="Save Report", command=lambda: save_report(report))
    save_button.pack(side=tk.LEFT, padx=10, pady=10)
    
    def close_window():
        window.destroy()
    window.protocol("WM_DELETE_WINDOW", close_window)  # Handle window close event
    window.mainloop()

def save_report(report):
    # Ask the user for the file path and type
    file_path = filedialog.asksaveasfilename(
        defaultextension=".txt",
        filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx")]
    )

    if not file_path:
        return  # User canceled the save dialog

    # Generate the report text
    report_text = "FILE AND LINK STATUS REPORT\n\n"
    
    for column, data in report.items():
        full_name = column_full_names.get(column, column)
        report_text += f"Column: {column} ({full_name})\n"
        report_text += f"  Total Count: {data['total_count']}\n"
        report_text += f"  'File Not Found' Count: {data['file_not_found_count']}\n"
        report_text += f"  'Link Not Found' Count: {data['link_not_found_count']}\n"
        
        # Add IDs for "File Not Found"
        if data['file_not_found_ids']:
            report_text += "\n  IDs with 'File Not Found':\n"
            for row in data['file_not_found_ids']:
                report_text += f"ID: {row[0]}, ID Tronc: {row[1]}, Code: {row[2]} \n"
        
        # Add IDs for "Link Not Found"
        if data['link_not_found_ids']:
            report_text += "\n  IDs with 'Link Not Found':\n"
            for row in data['link_not_found_ids']:
                report_text += f"ID: {row[0]}, ID Tronc: {row[1]}, Code: {row[2]} \n"
                
        report_text += "\n" + "-" * 80 + "\n\n"

    # Save the report
    if file_path.endswith(".txt"):
        with open(file_path, "w", encoding="utf-8") as file:
            file.write(report_text)
    elif file_path.endswith(".docx"):
        doc = Document()
        
        # Add title
        doc.add_heading("FILE AND LINK STATUS REPORT", level=1)
        
        for column, data in report.items():
            full_name = column_full_names.get(column, column)
            doc.add_heading(f"Column: {column} ({full_name})", level=2)
            
            # Add counts
            p = doc.add_paragraph()
            p.add_run(f"Total Count: {data['total_count']}\n")
            p.add_run(f"'File Not Found' Count: {data['file_not_found_count']}\n")
            p.add_run(f"'Link Not Found' Count: {data['link_not_found_count']}\n")
            
            # Add IDs for "File Not Found"
            if data['file_not_found_ids']:
                doc.add_heading("IDs with 'File Not Found':", level=3)
                for row in data['file_not_found_ids']:
                    doc.add_paragraph(f"ID: {row[0]}, ID Tronc: {row[1]}, Code: {row[2]}", style='List Bullet')
            
            # Add IDs for "Link Not Found"
            if data['link_not_found_ids']:
                doc.add_heading("IDs with 'Link Not Found':", level=3)
                for row in data['link_not_found_ids']:
                    doc.add_paragraph(f"ID: {row[0]}, ID Tronc: {row[1]}, Code: {row[2]}", style='List Bullet')
            
            doc.add_paragraph("-" * 80)
        
        doc.save(file_path)

    messagebox.showinfo("Success", f"Report saved to {file_path}")

def main(dbname, user, password, host, port, table_name):
    try:
        # Connect to the database
        conn = psycopg2.connect(
            dbname=dbname,
            user=user,
            password=password,
            host=host,
            port=port
        )
        # Generate the report
        report = generate_report(conn, table_name)
        # Display the report in a GUI
        display_report_gui(report)
        conn.close()
    except Exception as e:
        messagebox.showerror("Database Error", str(e))
        return

if __name__ == "__main__":
    main()
